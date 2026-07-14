# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document('docs/AI_거버넌스_지침_v3_0.docx')

def h1(doc, text): return doc.add_paragraph(text, style='Heading 1')
def h2(doc, text): return doc.add_paragraph(text, style='Heading 2')
def h3(doc, text): return doc.add_paragraph(text, style='Heading 3')
def body(doc, text): return doc.add_paragraph(text)
def li(doc, text): return doc.add_paragraph(text, style='List Paragraph')

new_paras = []

# 제9장
new_paras.append(h1(doc, '제9장 AX 위원회 및 실무 협의체 운영'))
new_paras.append(h2(doc, '제28조 (AX위원회 구성 및 운영)'))
new_paras.append(body(doc, '전사 AI 도입 전략의 최종 의사결정 기구로 AX위원회를 둔다.'))
new_paras.append(li(doc, '위원장은 AX/PI센터장이 맡으며, 위원은 CCO·CRO·CISO·현업임원 等으로 구성한다.'))
new_paras.append(li(doc, 'AX팀장을 간사로 하여 반기 1회 정기회의를 개최한다.'))
new_paras.append(h3(doc, 'AX위원회 의결 및 심의 사항'))
new_paras.append(li(doc, '전사 AI 도입 전략 및 연간 AX 과제 로드맵 승인'))
new_paras.append(li(doc, 'AI 도입 우선순위 결정 및 주요 AI 투자(예산) 승인'))
new_paras.append(li(doc, 'AI Agent의 효용성·안정성 최종 심의 및 Agent-HR 등재 승인'))
new_paras.append(li(doc, '우수 에이전트 포상 및 전사 확산 여부 결정'))

new_paras.append(h2(doc, '제29조 (AX 실무 협의체 구성 및 역할)'))
new_paras.append(body(doc, 'AX위원회 산하에 실무 협의체를 두어 Agent 운영 현황·개선·폐기를 실무 차원에서 관리한다.'))
new_paras.append(li(doc, '구성: 유관 부서장, AX팀장 / 운영: 월 1회 정기 및 비정기'))
new_paras.append(li(doc, '고위험 Agent 식별 및 통제 절차 수립'))
new_paras.append(li(doc, '현업 개발 Agent의 표준화 여부·안정성 검증 후 AX위원회 상정'))
new_paras.append(li(doc, 'Agent의 운영 현황, 개선방안, 폐기시점 협의 (Life-Cycle 관리)'))
new_paras.append(li(doc, 'AI 활용 결과 및 기여도를 분석해 Agent 성능평가 및 보상 제안'))

# 제10장
new_paras.append(h1(doc, '제10장 AI Agent 운영 관리'))
new_paras.append(h2(doc, '제30조 (Agent 위험 등급 분류)'))
new_paras.append(body(doc, 'AI Agent는 위험도 및 비즈니스 영향도에 따라 3개 등급으로 분류하며, 등급별 통제 수준을 달리 적용한다.'))
new_paras.append(li(doc, '고위험(High Risk): 대외 고객 대상 응대, 핵심 전략·개인정보 취급, 사람 개입 없이 업무 시스템 상태를 변경·실행하는 Agent. → 출시 전 보안·규제 심사 및 상시 모니터링 필수.'))
new_paras.append(li(doc, '중위험(Medium Risk): 내부 임직원 대상 업무 프로세스 자동화, 비기밀 데이터 기반 분석·기획서 초안 작성 등. → 최종 검토 시 Human-in-the-loop 적용.'))
new_paras.append(li(doc, '저위험(Low Risk): 단순 사내 규정 조회, 공개 사외 데이터 기반 리서치 지원 등. → 표준 운영 절차 준수.'))

new_paras.append(h2(doc, '제31조 (Agent 평가 및 보상)'))
new_paras.append(body(doc, 'AI 활용 결과 및 업무 기여도를 분기 1회 종합 평가하여 AX위원회에 보고한다.'))
new_paras.append(li(doc, '응답 정확도 및 업무 기여도를 정량·정성 평가한다.'))
new_paras.append(li(doc, "'우수' 및 '탁월' 에이전트를 분기별 선정하여 개발·운영 조직에 포상을 제안한다."))
new_paras.append(li(doc, '선정된 우수 Agent는 AX위원회 심의를 거쳐 전사 자산화 및 확산 여부를 결정한다.'))

new_paras.append(h2(doc, '제32조 (Context 자산화)'))
new_paras.append(body(doc, 'AI 활용 과정에서 생성된 프롬프트 템플릿, Fine-tuning 데이터셋, 산출물 등을 전사 지식 저장소에 등록·관리한다.'))
new_paras.append(li(doc, '산출물의 생성 주기 및 변경 이력을 추적 관리한다.'))
new_paras.append(li(doc, 'AI 모델 업데이트 및 사용자 변경에 따른 결과물 일관성을 확보한다.'))
new_paras.append(li(doc, '기밀등급에 따라 접근 권한을 분리·적용한다.'))

# 제11장
new_paras.append(h1(doc, '제11장 생성형 AI 도입 및 운영'))
new_paras.append(h2(doc, '제33조 (AI 리터러시 레벨별 도구 배분)'))
new_paras.append(body(doc, '업무 목적 및 비용 효율성을 고려하여 임직원의 AI 리터러시 수준에 따라 AI 도구를 배분한다.'))
new_paras.append(li(doc, 'Level 4 (Expert): AX센터·디지털마케팅·자산운용 부문 전문개발자 → AI연구망·AWS 샌드박스 (2026년 10월)'))
new_paras.append(li(doc, 'Level 3 (Advanced): AX크루 및 현업 코딩 가능자 50명 → Codex·Claude Code·Antigravity (2026년 10월)'))
new_paras.append(li(doc, 'Level 2 (Intensive): No/Low-Code Agent 개발 가능 현업 100명 → GPT Chat·Claude Cowork (2026년 8월)'))
new_paras.append(li(doc, 'Level 1 (Basic): 일반 임직원 100명 → Gemini·GPT for Excel (2026년 8월)'))

new_paras.append(h2(doc, '제34조 (모니터링 및 비용 관리)'))
new_paras.append(body(doc, '모델별 토큰(Token) 사용량·호출 빈도·응답속도 및 API 비용을 상시 모니터링한다.'))
new_paras.append(li(doc, '정기적인 비용 분석을 수행하고 경량 모델 대체 방안을 검토한다.'))
new_paras.append(li(doc, '관리자 환경에서 최대 토큰 사용량을 통제한다.'))

new_paras.append(h2(doc, '제35조 (AI 리터러시 제고)'))
new_paras.append(li(doc, '본부장·임원 교육은 인사팀 주관으로 진행한다.'))
new_paras.append(li(doc, '사내 강사를 육성하여 AI 크루 대상 강의를 진행한다.'))
new_paras.append(li(doc, '기업형 AI 도입·배분에 따른 사내교육을 추가 실시한다 (2026년 8월~).'))
new_paras.append(li(doc, 'AI 활용 우수 사례(Best Practice) 공모전을 개최하고 프롬프트 템플릿·가이드북을 배포한다.'))

# 제12장
new_paras.append(h1(doc, '제12장 PI T/F 지원 방안'))
new_paras.append(h2(doc, '제36조 (과제 수립 지원 및 PoC)'))
new_paras.append(body(doc, 'AX팀은 현업 부서의 프로세스를 분석하고 AI 적용 타당성을 검증한다.'))
new_paras.append(li(doc, '요구 데이터 검토(보유·품질·권한) 및 프로토타입 개발을 지원한다.'))
new_paras.append(li(doc, '정량적·정성적 성공 지표를 수립하고 검증 및 상용화 판정을 수행한다.'))

new_paras.append(h2(doc, '제37조 (보안 등급 기반 단계적 실행)'))
new_paras.append(h3(doc, 'Phase 1 — 외부 데이터 기반 부서 우선 적용'))
new_paras.append(li(doc, '대상: 상품기획·마케팅·운용 등 시장 동향 리서치 및 외부 공개 데이터 활용 비중이 높은 부서'))
new_paras.append(li(doc, '클라우드 기반 AI 인프라를 활용하여 성공 사례를 확보한다.'))
new_paras.append(h3(doc, 'Phase 2 — 내부 기밀·대외비 데이터 취급 부서 (2026년 10월~)'))
new_paras.append(li(doc, '대상: 경영지원·리스크·컴플라이언스 등 보안 등급이 높은 부서'))
new_paras.append(li(doc, '구축 중인 내부(On-premise) AI 플랫폼과 연계하여 진행한다.'))

new_paras.append(h2(doc, '제38조 (데이터 수급 및 인프라 연계)'))
new_paras.append(li(doc, '스노우플레이크(Snowflake) 기반 전용 데이터 마트(Data Mart)를 구축한다.'))
new_paras.append(li(doc, 'MCP(Model Context Protocol) 서버 커넥터를 구축하여 LLM이 데이터를 직접 참조할 수 있는 환경을 제공한다.'))
new_paras.append(li(doc, '데이터의 정제·라벨링·가명정보 처리 표준 가이드라인을 제공하여 AI Ready 데이터를 확보한다.'))

# 부칙 단락 앞에 삽입
buching_para = doc.paragraphs[65]
for np in new_paras:
    buching_para._p.addprevious(np._p)

# 버전 메타 업데이트
for para in doc.paragraphs:
    if 'AX-POLICY-2026-001' in para.text and 'v3.0' in para.text:
        para.text = 'AX-POLICY-2026-001  v4.0'
    elif '개정일' in para.text and 'v3.0' in para.text:
        para.text = '제정일: 2026년 7월 2일  |  개정일: 2026년 7월 8일 (v4.0 — AX위원회·Agent운영관리·리터러시·PI T/F 추가)'
    elif 'v3.0 추가 내용' in para.text:
        para.text = '[v4.0 추가 내용] 제9장(AX위원회·실무협의체), 제10장(AI Agent 운영관리), 제11장(생성형AI 도입·운영), 제12장(PI T/F 지원방안)이 신설되었습니다. 제1~8장은 v3.0과 동일합니다.'
    elif '제2조 (기존 지침과의 관계)' in para.text:
        para.text = '제2조 (기존 지침과의 관계)  제1~8장 및 기존 부칙은 v3.0과 동일하다. v4.0은 제9~12장을 추가한다.'
    elif '제1조 (시행일)' in para.text and 'v3.0' in para.text:
        para.text = '제1조 (시행일)  이 지침(v4.0)은 2026년 7월 8일부터 시행한다.'
    elif 'v3.0 (2026.07.05' in para.text:
        new_hist = doc.add_paragraph('v4.0 (2026.07.08 — AX위원회, Agent 운영관리, 리터러시 레벨, PI T/F 지원 추가)', style='List Paragraph')
        para._p.addnext(new_hist._p)

doc.save('docs/AI_거버넌스_지침_v4_0.docx')
print('완료: AI_거버넌스_지침_v4_0.docx 저장됨')
