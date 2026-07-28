import sys
sys.stdout.reconfigure(encoding="utf-8")
from docx import Document

doc = Document("docs/AX-REGULATION-2026-001_AI운영규정.docx")

def get_style(name):
    for s in doc.styles:
        if s.name == name:
            return s
    return None

h1 = get_style("Heading 1")
h2 = get_style("Heading 2")
li = get_style("List Paragraph")

def p(text, style=None):
    para = doc.add_paragraph(text)
    if style: para.style = style
    return para

new_paras = []
new_paras.append(p("제6장  데이터 통제 방안", h1))
new_paras.append(p("제11조 (데이터 유출 방지)", h2))
new_paras.append(p("회사는 AI 서비스 활용 과정에서 발생할 수 있는 데이터 유출을 방지하기 위해 다음 통제를 적용한다."))
new_paras.append(p("DLP(Data Loss Prevention) 솔루션을 AI 서비스 접점에 연동하여 기밀데이터 외부 전송을 실시간 차단한다.", li))
new_paras.append(p("G2 내부, G3 기밀 데이터는 전송 전 마스킹 또는 익명화 처리를 의무화한다.", li))
new_paras.append(p("AI 서비스에 입력된 데이터의 로그는 90일간 보존하고 분기 1회 감사팀이 점검한다.", li))
new_paras.append(p("제12조 (암호화 기준)", h2))
new_paras.append(p("AI 관련 데이터 저장·전송 시 다음 암호화 기준을 준수한다."))
new_paras.append(p("전송 구간: TLS 1.2 이상 적용. 외부 클라우드 API 호출 시 HTTPS 필수.", li))
new_paras.append(p("저장 구간: G2 이상 데이터는 AES-256 이상으로 암호화하여 저장한다.", li))
new_paras.append(p("키 관리: 암호화 키는 데이터와 분리된 키 관리 시스템(KMS)에서 관리한다.", li))
new_paras.append(p("제13조 (보안 도구 사용 제한)", h2))
new_paras.append(p("AX팀이 승인하지 않은 외부 AI 보조 개발 도구(Copilot, Cursor 등)를 업무용 PC 및 사내 시스템에 연동하는 것을 금지한다."))
new_paras.append(p("단, AX팀장이 보안 검토 후 승인한 도구는 제한적으로 허용하며, 사용 현황을 분기 1회 보고한다.", li))
new_paras.append(p("승인 도구라도 G3 기밀 데이터를 해당 도구에 입력하거나 연동하는 것은 금지한다.", li))
new_paras.append(p("제7장  데이터 보존 및 파기", h1))
new_paras.append(p("제14조 (데이터 보존 기간)", h2))
new_paras.append(p("AI 서비스 활용 과정에서 생성·처리된 데이터는 관련 법령 및 내부 규정에 따라 다음 기간 동안 보존한다."))
new_paras.append(p("투자권유 관련 AI 활용 기록 (프롬프트, 생성 결과물 포함): 10년", li))
new_paras.append(p("금융상품 판매·운용 관련 AI 생성 분석 자료: 15년", li))
new_paras.append(p("개인신용정보가 포함된 AI 처리 데이터: 이용 목적 달성 후 즉시 파기 또는 최대 5년 이내", li))
new_paras.append(p("일반 업무 효율화 관련 AI 생성 산출물: 3년", li))
new_paras.append(p("제15조 (데이터 파기 절차)", h2))
new_paras.append(p("보존 기간이 경과한 데이터는 다음 절차에 따라 파기한다."))
new_paras.append(p("파기 대상 목록을 반기 1회 작성하고 데이터플랫폼팀 및 컴플라이언스팀의 확인을 받는다.", li))
new_paras.append(p("전자적 데이터는 복구 불가능한 방법(NIST 800-88 준거)으로 삭제한다.", li))
new_paras.append(p("파기 결과는 파기 일자·방법·담당자를 포함한 파기 확인서를 작성하여 3년간 보존한다.", li))
new_paras.append(p("외부 AI 서비스 사업자에게 전송된 데이터의 파기 요청 절차는 AX팀이 각 사업자의 정책에 따라 처리한다.", li))

buching_para = None
for para in doc.paragraphs:
    if "부  칙" in para.text or "부칙" in para.text:
        buching_para = para
        break

if buching_para:
    for np in new_paras:
        buching_para._p.addprevious(np._p)
    print("삽입 완료")
else:
    print("부칙 단락 못 찾음")

# 개정일 업데이트
for para in doc.paragraphs:
    if "AX-REGULATION-2026-001" in para.text and "2026-07" in para.text:
        para.text = "AX-REGULATION-2026-001  |  v2.0  |  개정 2026-07-08"
        break
    elif "AX-REGULATION-2026-001" in para.text:
        para.text = "AX-REGULATION-2026-001  |  v2.0  |  개정 2026-07-08"
        break

for para in doc.paragraphs:
    if "제1조 (시행일)" in para.text and "2026년 7월 2일" in para.text:
        para.text = "제1조 (시행일)  이 규정은 2026년 7월 2일부터 시행한다. (v2.0 개정 2026년 7월 8일)"
        break

doc.save("docs/AX-REGULATION-2026-001_AI운영규정.docx")
print("저장 완료")